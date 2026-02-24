"""
FileProcessor — извлечение текста из PDF, DOCX, TXT и валидация загруженных файлов.

Пайплайн:
  Файл → проверка расширения → проверка размера → извлечение текста →
  проверка кодировки → проверка текстового слоя (PDF) → подсчёт слов →
  валидация объёма → возврат extracted_text + метаданные
"""

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Result data classes
# ---------------------------------------------------------------------------


@dataclass
class ValidationResult:
    """Результат валидации файла."""

    valid: bool = True
    error_code: Optional[str] = None
    error_message: Optional[str] = None
    word_count: int = 0
    warnings: List[str] = field(default_factory=list)


@dataclass
class ExtractionResult:
    """Результат извлечения текста из файла."""

    ok: bool = True
    extracted_text: str = ""
    word_count: int = 0
    file_info: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    error_code: Optional[str] = None
    error_message: Optional[str] = None


# ---------------------------------------------------------------------------
# FileProcessor
# ---------------------------------------------------------------------------


class FileProcessor:
    """Обработчик загруженных файлов: валидация + извлечение текста."""

    DEFAULT_ALLOWED_EXTENSIONS = [".pdf", ".docx", ".txt"]
    DEFAULT_MAX_FILE_SIZE_MB = 18
    DEFAULT_MAX_WORD_COUNT = 15000
    DEFAULT_MIN_WORD_COUNT = 50

    def __init__(
        self,
        allowed_extensions: Optional[List[str]] = None,
        max_file_size_mb: int = DEFAULT_MAX_FILE_SIZE_MB,
        max_word_count: int = DEFAULT_MAX_WORD_COUNT,
        min_word_count: int = DEFAULT_MIN_WORD_COUNT,
    ):
        self.allowed_extensions = allowed_extensions or self.DEFAULT_ALLOWED_EXTENSIONS
        self.max_file_size_mb = max_file_size_mb
        self.max_word_count = max_word_count
        self.min_word_count = min_word_count

    # ------------------------------------------------------------------
    # Main entry point
    # ------------------------------------------------------------------

    def process_file(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """
        Полный пайплайн: валидация → извлечение текста → проверка объёма.

        Args:
            file_bytes: Содержимое файла в байтах.
            filename: Оригинальное имя файла (нужно для определения расширения).

        Returns:
            ExtractionResult с текстом или ошибкой.
        """
        ext = Path(filename).suffix.lower()

        # 1. Проверка расширения
        if ext not in self.allowed_extensions:
            return ExtractionResult(
                ok=False,
                error_code="unsupported_format",
                error_message=f"Формат файла {ext} не поддерживается. "
                              f"Поддерживаемые форматы: {', '.join(e.upper().lstrip('.') for e in self.allowed_extensions)}.",
            )

        # 2. Проверка размера
        file_size_mb = len(file_bytes) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            return ExtractionResult(
                ok=False,
                error_code="file_too_large",
                error_message=f"Файл слишком большой ({file_size_mb:.1f} МБ). "
                              f"Максимальный размер — {self.max_file_size_mb} МБ. "
                              "Попробуйте загрузить отдельные главы.",
            )

        # 3. Проверка на пустой файл
        if len(file_bytes) == 0:
            return ExtractionResult(
                ok=False,
                error_code="file_empty",
                error_message="Файл пуст. Проверьте, что вы загрузили правильный файл.",
            )

        # 4. Извлечение текста
        file_info: Dict[str, Any] = {
            "original_name": filename,
            "format": ext.lstrip("."),
            "size_mb": round(file_size_mb, 2),
        }
        warnings: List[str] = []

        try:
            if ext == ".txt":
                text, encoding_used = self.extract_text_from_txt(file_bytes)
                file_info["encoding"] = encoding_used
            elif ext == ".pdf":
                text, has_text_layer, pages = self.extract_text_from_pdf(file_bytes)
                file_info["has_text_layer"] = has_text_layer
                file_info["pages"] = pages
                if not has_text_layer or not text.strip():
                    return ExtractionResult(
                        ok=False,
                        error_code="no_text_layer",
                        error_message="В этом PDF-файле нет текста — он содержит только изображения "
                                      "(например, отсканированные страницы). Чтобы мы могли обработать "
                                      "материал, нужен файл с «настоящим» текстом, который можно выделить "
                                      "и скопировать. Попробуйте найти электронную версию учебника в "
                                      "формате DOCX или PDF с текстом.",
                        file_info=file_info,
                    )
            elif ext == ".docx":
                text = self.extract_text_from_docx(file_bytes)
            else:
                return ExtractionResult(
                    ok=False,
                    error_code="unsupported_format",
                    error_message=f"Формат {ext} не поддерживается.",
                )
        except ImportError as ie:
            logger.error("[FileProcessor] Missing library: %s", ie)
            return ExtractionResult(
                ok=False,
                error_code="server_missing_library",
                error_message="Серверу не хватает библиотеки для обработки этого формата файла. "
                              "Обратитесь к администратору.",
            )
        except Exception as exc:
            logger.exception("[FileProcessor] Text extraction failed: %s", exc)
            return ExtractionResult(
                ok=False,
                error_code="extraction_failed",
                error_message="Не удалось корректно прочитать текст из файла. "
                              "Попробуйте сохранить его в формате DOCX или UTF-8 TXT.",
            )

        # 5. Валидация объёма текста
        validation = self.validate_text(text)
        file_info["word_count"] = validation.word_count
        warnings.extend(validation.warnings)

        if not validation.valid:
            return ExtractionResult(
                ok=False,
                error_code=validation.error_code,
                error_message=validation.error_message,
                word_count=validation.word_count,
                file_info=file_info,
                warnings=warnings,
            )

        return ExtractionResult(
            ok=True,
            extracted_text=text,
            word_count=validation.word_count,
            file_info=file_info,
            warnings=warnings,
        )

    # ------------------------------------------------------------------
    # Text validation
    # ------------------------------------------------------------------

    def validate_text(self, text: str) -> ValidationResult:
        """Валидация извлечённого текста по объёму."""
        word_count = len(text.split())
        warnings: List[str] = []

        if word_count < self.min_word_count:
            return ValidationResult(
                valid=False,
                error_code="too_few_words",
                error_message="В файле слишком мало текста для генерации заданий. "
                              f"Нужно хотя бы {self.min_word_count} слов учебного материала.",
                word_count=word_count,
            )

        if word_count > self.max_word_count:
            warnings.append(
                f"Материал очень объёмный (~{word_count} слов, ~{word_count // 250} стр.). "
                "Рекомендуем загрузить одну-две главы."
            )

        return ValidationResult(
            valid=True,
            word_count=word_count,
            warnings=warnings,
        )

    def validate_file(self, file_bytes: bytes, filename: str) -> ValidationResult:
        """Лёгкая валидация файла без извлечения текста (для фронтенда)."""
        ext = Path(filename).suffix.lower()

        if ext not in self.allowed_extensions:
            return ValidationResult(
                valid=False,
                error_code="unsupported_format",
                error_message=f"Формат файла {ext} не поддерживается. "
                              f"Поддерживаемые: {', '.join(e.upper().lstrip('.') for e in self.allowed_extensions)}.",
            )

        file_size_mb = len(file_bytes) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            return ValidationResult(
                valid=False,
                error_code="file_too_large",
                error_message=f"Файл слишком большой ({file_size_mb:.1f} МБ). "
                              f"Максимум — {self.max_file_size_mb} МБ.",
            )

        if len(file_bytes) == 0:
            return ValidationResult(
                valid=False,
                error_code="file_empty",
                error_message="Файл пуст.",
            )

        return ValidationResult(valid=True)

    # ------------------------------------------------------------------
    # Text extraction: TXT
    # ------------------------------------------------------------------

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> Tuple[str, str]:
        """
        Извлечь текст из TXT с определением кодировки.

        Returns:
            (text, encoding_used)
        """
        # 1. UTF-8
        try:
            text = file_bytes.decode("utf-8")
            # Убираем BOM если есть
            if text.startswith("\ufeff"):
                text = text[1:]
            return text, "utf-8"
        except UnicodeDecodeError:
            pass

        # 2. chardet
        try:
            import chardet
            detected = chardet.detect(file_bytes)
            encoding = detected.get("encoding")
            confidence = detected.get("confidence", 0)
            if confidence >= 0.7 and encoding:
                try:
                    return file_bytes.decode(encoding), encoding
                except (UnicodeDecodeError, LookupError):
                    pass
        except ImportError:
            logger.debug("[FileProcessor] chardet not available, skipping detection")

        # 3. Windows-1251 (частая для русскоязычных документов)
        try:
            return file_bytes.decode("cp1251"), "cp1251"
        except UnicodeDecodeError:
            pass

        # 4. KOI8-R
        try:
            return file_bytes.decode("koi8-r"), "koi8-r"
        except UnicodeDecodeError:
            pass

        # 5. Крайний случай
        return file_bytes.decode("utf-8", errors="replace"), "utf-8 (lossy)"

    # ------------------------------------------------------------------
    # Text extraction: PDF
    # ------------------------------------------------------------------

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, bool, int]:
        """
        Извлечь текст из PDF.

        Returns:
            (text, has_text_layer, page_count)
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF (fitz) is required for PDF processing. "
                "Install: pip install PyMuPDF"
            )

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = doc.page_count
        text_parts: List[str] = []
        has_text = False
        image_heavy_pages = 0

        for page in doc:
            page_text = page.get_text("text").strip()
            if page_text:
                has_text = True
                text_parts.append(page_text)
            else:
                image_heavy_pages += 1

        doc.close()

        text = "\n\n".join(text_parts)

        return text, has_text, pages

    @staticmethod
    def has_text_layer(file_bytes: bytes) -> bool:
        """Проверка наличия текстового слоя в PDF (без полного извлечения)."""
        try:
            import fitz
        except ImportError:
            return False

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        has_text = False
        # Проверяем первые 3 страницы
        for i, page in enumerate(doc):
            if i >= 3:
                break
            if page.get_text("text").strip():
                has_text = True
                break
        doc.close()
        return has_text

    # ------------------------------------------------------------------
    # Text extraction: DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Извлечь текст из DOCX (параграфы + таблицы)."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install: pip install python-docx"
            )

        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []

        # Параграфы
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Таблицы
        seen_cells: set = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen_cells:
                        seen_cells.add(cell_text)
                        parts.append(cell_text)

        return "\n".join(parts)

    # ------------------------------------------------------------------
    # Encoding detection utility
    # ------------------------------------------------------------------

    @staticmethod
    def detect_encoding(file_bytes: bytes) -> str:
        """
        Определить кодировку файла.

        Returns:
            Название кодировки.
        """
        # UTF-8 BOM
        if file_bytes[:3] == b"\xef\xbb\xbf":
            return "utf-8-sig"

        # Попытка UTF-8
        try:
            file_bytes.decode("utf-8")
            return "utf-8"
        except UnicodeDecodeError:
            pass

        # chardet
        try:
            import chardet
            detected = chardet.detect(file_bytes)
            encoding = detected.get("encoding")
            confidence = detected.get("confidence", 0)
            if confidence >= 0.7 and encoding:
                return encoding
        except ImportError:
            pass

        # Эвристика для кириллицы: проверяем cp1251
        try:
            file_bytes.decode("cp1251")
            return "cp1251"
        except UnicodeDecodeError:
            pass

        return "unknown"
